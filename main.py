from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from typing import List
from docx import Document as WordDoc
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import openai
import os
import asyncio
import re
import json
from pathlib import Path
from fastapi import Request
from fastapi.responses import JSONResponse
import datetime
from bs4 import BeautifulSoup

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://merilmoretolife.github.io"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

openai.api_key = os.getenv("OPENAI_API_KEY")

# --- Data Models ---
class DeviceRequest(BaseModel):
    deviceName: str
    intendedUse: str
    sections: list[str]

class FinalizedDevice(BaseModel):
    deviceName: str
    intendedUse: str
    designInputHtml: str
    finalizedBy: str
    diComplete: bool
    doComplete: bool
    finalizedAt: str
    sections: list[str]  # ✅ Add this line

class DesignOutputRequest(BaseModel):
    deviceName: str
    intendedUse: str
    section: str

class UpdateRequest(BaseModel):
    deviceName: str
    intendedUse: str
    section: str
    currentContent: str
    remark: str

# --- Helpers ---
def insert_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_prompt(device_name: str, intended_use: str, section: str) -> str:
    intro = f"Generate design input content for the medical device '{device_name}', intended for '{intended_use}', under the section: '{section}'. Please follow the specified format, adjust details as per the device type and Use globally accepted medtech regulatory language.\n\n"

    instructions = {
        "Functional and Performance Requirements": f"""
Include the following:
1. Material of Construction – List main materials and cite relevant ASTM/ISO standards based on the device type.
2. Component Design and Dimension – Define critical design features and tolerances.
3. Mechanical Properties and tests with relevant/applicable USP/ISO/ASTM,etc. standards.
Tailor content based on whether the device is an implant, instrument, or external device.
""",

        "Biological and Safety Requirements": f"""
Include:
1. Raw Material Compatibility – Mention inertness, sterilization tolerance, and chemical compatibility.
2. Biological Safety – Address cytotoxicity, irritation, sensitization, systemic effects.
3. Biocompatibility Tests Required – Based on ISO 10993-1 and contact duration, list applicable tests from ISO 10993 series and USP <87>/<88>.
4. Applicable Standards – List ISO 10993-1 and USP references only here.
Base all content on the nature, location, and duration of contact.
""",

        "Labeling and IFU Requirements": f"""
Include:
1. Label Information – List key product identifiers (name, code, lot, expiry, symbols).
2. Labeling Standards – Mention EN ISO 15223-1, EN ISO 20417, 21 CFR 801.109, ISO 14630.
3. IFU and e-IFU Requirements – Include indication, warnings, usage, multilingual needs, and digital/e-IFU compliance under EU Regulation 207/2012.
Tailor label/IFU fields based on region and class.
""",

        "Sterilization Requirements": f"""
Include:
1. Sterilization Method – Recommend EO, Steam, Gamma, or others based on material.
2. Applicable Standards – ISO 11135, ISO 11137, ISO 17665, ISO 10993-7, ISO 11737-1/2, USP <71>, <85>, <61>.
3. Required Tests – Bioburden, SAL 10⁻⁶, residuals, endotoxins, seal integrity.
Ensure compatibility with device sensitivity and configuration.
""",

        "Stability / Shelf Life Requirements": f"""
Include:
1. Shelf Life Objective – Specify target duration based on comparable products.
2. Factors Impacting Stability – Temperature, humidity, UV exposure, packaging.
3. Stability Study – Real-time and accelerated aging (ASTM F1980), post-aging validation.
4. Applicable Standards – ASTM F1980, ISO 11607-1, ICH Q1A(R2).
""",

        "Packaging and Shipping Requirements": f"""
Include:
1. Packaging Objectives – Protect from light, moisture, contamination, damage, maintain sterility.
2. Packaging Materials – Detail materials used (Tyvek, foil, blister), and barrier properties.
3. Packaging Configuration – Describe primary, secondary, tertiary setup and inclusion of IFU.
4. Standards – EN ISO 11607-1/2, ASTM F88.
Adjust for product fragility, sterility, and logistics.
""",

        "Manufacturing Requirements": f"""
Include:
1. Facility Infrastructure – GMP design: epoxy flooring, HEPA filters, material finishes.
2. Cleanroom Classification – ISO Class 8 or better depending on operation.
3. Equipment and Sanitation – GMP-compliant equipment, cleaning protocols.
4. QC and Storage – Environmental control, microbiology lab, USP testing capability.
Standards: ISO 13485, 21 CFR Part 820.
""",

        "Statutory and Regulatory Requirements": f"""
Include:
1. Indian Regulatory – CDSCO rules, classification (A–D), MD-13, MD-9, ISO 13485.
2. EU Regulatory – CE Marking, Detailed EU MDR classification (Class and Rule) from https://eur-lex.europa.eu/legal-content/EN/TXT/PDF/?uri=CELEX:32017R0745, GSPR, Technical File, ISO 13485.
3. US FDA – Class I/II/III, 510(k)/PMA, QSR (21 CFR Part 820), Establishment Registration.
Tailor classification and pathways based on device use and risk.
"""
    }

    return intro + instructions.get(section, "Provide general content.")

def generate_do_prompt(device_name: str, intended_use: str, section: str) -> str:
    if section == "Functional and Performance Requirements":
        return f"""
Generate the Design Output for a medical device called '{device_name}', intended for '{intended_use}', under the section: 'Functional and Performance Requirements'.

Include the following clearly, using clean formatting and tables where applicable:

1. Material of Construction:
- Specify exact materials, and cite relevant ISO/ASTM standards used for material validation in reference to its design input.

2. Component Design and Dimension:
- Define dimensional requirements and allowable tolerances, preferably in table format.
- Highlight size range if applicable.

3. Mechanical Properties:
- Applicable Mechanical Properties and Tests conducted and expected limits or acceptance criteria.
- Mention applicable standards - USP/ISO/ASTM.

Consider Design Input for giving Output Results.
Base the output on relevant real standards like USP, ISO, ASTM. Include tables with actual parameter ranges (e.g., tensile strength by USP size). Avoid generalizations.
"""

    elif section == "Biological and Safety Requirements":
        return f"""
Generate the Design Output for a medical device called '{device_name}', intended for '{intended_use}', under the section: 'Biological and Safety Requirements'.

Only include one subsection:

## 1. Biocompatibility Tests Requirements

- Inject this statement every time before table: "Based on the nature and duration of body contact, following is the list of all required biocompatibility tests as per ISO 10993-1."
- Format the information as a markdown table with the following columns:

| Sr. No. | Standard Reference | Study Name | Study No. |
|---------|--------------------|------------|-----------|

- Include standard numbers (e.g., ISO 10993-5:2009, USP <87>, USP <88>, etc...).
- Leave the "Study No." column blank.
- Do not include any notes or extra text outside the table.
"""

    elif section == "Packaging and Shipping Requirements":
        return f"""
Generate the Design Output for a medical device called '{device_name}', intended for '{intended_use}', under the section: 'Packaging and Shipping Requirements'.

Include a clear, product-specific packaging configuration (e.g., for surgical sutures). Mention:

1. Primary Packaging: e.g., suture wound in 8-shape, in paper/plastic tray, etc.
2. Secondary Packaging: pouch (e.g., aluminum), and box with IFU.
3. Sterility Maintenance: mention compatibility with sterilization and shelf life.
4. Qualification Tests: include the table below based on ASTM and USP standards.
5. Transportation Tests: include all relevant ASTM and IS references.

### Packaging Qualification Tests

| Parameter        | Acceptance Criteria  |
|------------------|----------------------|
| Seal Strength    | ≥ 2N                 |
| Seal Width       | ≥ 5mm                |
| Seal Integrity   | No Leakage           |
| Sterility        | USP <71>             |

### Transportation Tests

Mention the following:
- ASTM D4169-16: Performance Testing of Shipping Containers and Systems
- ASTM D5276: Drop Test of Loaded Containers by Free Fall
- ASTM D999: Vibration Testing of Shipping Containers
- IS 7028-4: Vertical Impact Drop Test
- IS 7028-2: Vibration Test at Fixed Low Frequency

The packaging configuration must ensure maintenance of sterility, physical integrity, and resistance during transport.
"""

    elif section == "Labeling and IFU Requirements":
        return f"""
Generate the Design Output for the medical device '{device_name}', intended for '{intended_use}', under the section: 'Labeling and IFU Requirements'.

1. **Applicable Labeling Standards**  
Mention the use of the following standards and regulations:
- EN ISO 15223-1: Medical devices — Symbols to be used with medical device labels, labeling, and information to be supplied  
- ISO 20417:2021: Medical devices — Information to be supplied by the manufacturer  
- 21 CFR Part 801: Labeling requirements by US FDA  
- Regulation (EU) 2021/2226: Requirements for electronic Instructions for Use (e-IFU)  

2. **Labeling Strategy**  
Explain labeling on:
- **Primary Pack** (e.g., Tyvek lid, blister, pouch)
- **Secondary Pack** (e.g., carton or box)
- **IFU** (paper and/or e-IFU)
Clarify symbols used, regional considerations, and where each type of information will appear.

3. **Labeling Content Table**  
Include a sample markdown table like the one below (add/remove rows based on the device type). Use ✓ for applicable, X for not applicable.

| Sr. No. | Labelling Requirement       | Primary Pack | Secondary Pack | IFU |
|---------|------------------------------|--------------|----------------|-----|
| 1       | Proprietary name of device   | ✓            | ✓              | ✓   |
| 2       | Description of the device    | ✓            | ✓              | ✓   |
| 3       | Intended Use                 | X            | ✓              | ✓   |
| 4       | Storage Conditions           | X            | ✓              | ✓   |
| 5       | Sterilization Method         | ✓            | ✓              | ✓   |
| ...     | ...                          | ...          | ...            | ... |

⚠️ *Note: The above table is illustrative. Actual fields must be tailored per device category, risk class, and market-specific requirements.*

4. **e-IFU Compliance**  
- Mention if e-IFU is applicable and the conditions under Regulation (EU) 2021/2226.
- Describe the access method (e.g., QR code, website), and ensure redundancy in case of digital access failure.
"""

    elif section == "Sterilization Requirements":
        return f"""
Generate the Design Output for the medical device '{device_name}', intended for '{intended_use}', under the section: 'Sterilization Requirements'.

Tailor the output based on the device’s nature, material, and packaging. Mention selected sterilization method(s), applicable standards, test requirements, and acceptance criteria. Also include bioburden, endotoxins, and residuals where applicable.

Include the following:

1. **Selected Sterilization Method(s)**  
- Clearly state the primary and (if applicable) secondary sterilization methods used for the device.  
- Example methods: Gamma Irradiation, Ethylene Oxide (EO), Steam, Dry Heat, etc.

2. **Applicable Standards for Sterilization**  
List only relevant standards (e.g.,):  
- EN ISO 11135:2014 (Ethylene Oxide)  
- EN ISO 11737-1 & 11737-2 (Bioburden & Sterility Testing)  
- EN ISO 20857 (Dry Heat)  
- EN ISO 11137-1/2 (Gamma)  
- USP <71>, <85>, <61>, <62>  

3. **Bioburden Test Requirements**  
Include a table like the one below if bioburden is applicable:

| Parameter                  | Acceptance Criteria      |
|----------------------------|--------------------------|
| Total Aerobic Viable Count | ≤ 1000 cfu/sample        |
| Total Fungal Count         | ≤ 100 cfu/sample         |

4. **Sterility Test / SAL**  
- Sterility assurance level must comply with a minimum SAL of 10⁻⁶.  
- Sterility to be confirmed as per USP <71> or ISO 11737-2.  

5. **Bacterial Endotoxin Limits**  
- State if applicable:  
  Example: Bacterial endotoxin level must not exceed 10 EU/device, as per USP <85>.

6. **Residuals (for EO sterilized devices)**  
If EO sterilization is used, include a table:

| Residual Component       | Maximum Limit (per device) |
|--------------------------|----------------------------|
| Ethylene Oxide (EO)      | ≤ 4 mg                     |
| Ethylene Chlorhydrin     | ≤ 9 mg                     |
| Ethylene Glycol (EG)     | ≤ 9 mg                     |

Explain that these limits are per ISO 10993-7 for EO residuals.

Summarize how the sterilization approach aligns with device material, intended use, and packaging configuration.
"""

    elif section == "Stability / Shelf Life Requirements":
        return f"""
Generate the Design Output for the medical device '{device_name}', intended for '{intended_use}', under the section: 'Stability / Shelf Life Requirements'.

Focus on how the shelf life of the device is validated and established. Tailor the output based on product type (sterile, implant, suture, etc.).

Include the following:

1. **Study References and Guidelines**  
- ICH Q1A(R2): Stability Testing of New Drug Substances and Products  
- ASTM F1980: Guide for Accelerated Aging of Sterile Barrier Systems  
- ISO 11607-1: Packaging for Terminally Sterilized Medical Devices

2. **Accelerated Aging Study**  
- State test conditions: 50°C ± 2°C and 75% RH ± 5% RH  
- Duration of study and intervals for testing (e.g., 1, 2, 3 months)  
- Acceptance criteria: critical parameters (e.g., sterility, tensile strength, packaging seal integrity) must remain within limits

3. **Real-Time Shelf Life Study**  
- Storage conditions: 30°C ± 2°C and 65% RH ± 5% RH  
- State intervals for evaluation (e.g., 3, 6, 9, 12 months)  
- Confirm that packaging and device performance are monitored

4. **Establishing Shelf Life**  
- Based on validated results of accelerated and real-time data  
- Final assigned shelf life (in months/years)  
- Mention if the shelf life is applicable to both packaging and device

Conclude how the validated stability studies justify the claimed shelf life in the labeling and regulatory documents.
"""


    elif section == "Manufacturing Requirements":
        return f"""
Generate the Design Output for the medical device '{device_name}', intended for '{intended_use}', under the section: 'Manufacturing Requirements'.

Base the output on applicable design input and general good manufacturing practices, especially for cleanroom-class devices or sterile implants. Consider variations based on device type.

Include the following:

1. **Facility Infrastructure and Layout**  
- Describe requirements for the premises including wall, floor, ceiling finishes (impervious, epoxy-coated, non-flaking).  
- Reference design elements such as coving, GMP zoning, and drainless layouts where applicable.  

2. **Cleanroom Classification and Environment Control**  
- Specify cleanroom ISO classes based on the device’s exposure and critical operations (e.g., ISO Class 8 or better for sterile steps).  
- Detail HVAC/air handling systems, number of air changes per hour, pressure differentials, HEPA filter use.

3. **Equipment and Utilities**  
- List critical equipment for production and testing (e.g., melting, molding, sealing for bone wax; injection molding for polymer; laser welding, etc.).  
- Ensure GMP compliance of equipment (e.g., SS316 construction, cleanability, calibration, and validation needs).

4. **Sanitation and Housekeeping**  
- Describe cleaning schedules, disinfectant rotation, cleaning validation if applicable, and documentation control.  
- Mention pest control, gowning procedures, and personnel hygiene measures.

5. **Storage and Material Handling**  
- Outline raw material and finished goods storage requirements — include temperature/humidity control, segregation, FIFO/FEFO logic.

6. **Quality Control (QC) and Microbiological Testing**  
- Describe in-house QC labs, including capability to test incoming raw materials, in-process samples, and finished goods.  
- Mention microbiological monitoring, if the product requires it, including environmental monitoring and bioburden/endotoxin tests.

7. **Regulatory Compliance**  
- The facility must operate under a Quality Management System compliant with ISO 13485, 21 CFR Part 820, and applicable local regulations.  

Emphasize how the manufacturing infrastructure aligns with design input and the intended use of the device. Tailor for device-specific considerations.
"""


    elif section == "Statutory and Regulatory Requirements":
        return f"""
Generate the Design Output for the medical device '{device_name}', intended for '{intended_use}', under the section: 'Statutory and Regulatory Requirements'.

Summarize applicable regulatory requirements in the following structure:


## 1. Indian Regulatory Requirements

Include:
- Manufacturing License (Form MD-9 or MD-5 depending on class)
- Factory License
- Quality Management System (ISO 13485:2016)

### Table: Indian Regulatory Compliance

| Sr. No. | Requirements             | Source / Guideline                                                                 | Process Description            |
|---------|--------------------------|-------------------------------------------------------------------------------------|--------------------------------|
| 1       | Manufacturing License (Form MD-9 or Form MD-5) | CDSCO Online Portal: https://cdscomdonline.gov.in/NewMedDev/Homepage               | Application & Approval         |
| 2       | Factory License          | https://dish.gujarat.gov.in/new-factory-license-application.htm                     | Factory Setup Compliance       |
| 3       | ISO 13485:2016 QMS       | ISO 13485:2016                                                                      | QMS Documentation & Certification |



## 2. European Union – CE Marking

Summarize CE regulatory pathway under EU MDR (Regulation (EU) 2017/745). Highlight:

- CE Certification
- Technical File development
- Conformity assessment route

### Table: EU MDR Compliance

| Sr. No. | Requirement      | Source / Regulation                           | Process Description        |
|---------|------------------|-----------------------------------------------|----------------------------|
| 1       | CE Certification | Regulation (EU) 2017/745 – EU MDR              | Technical File Preparation, Notified Body Involvement |


## 3. United States – US FDA

Mention applicable US FDA pathway (510(k), PMA, or Exempt) and QSR compliance.

### Table: USFDA Compliance

| Sr. No. | Requirement | Source / Guideline                        | Process Description            |
|---------|-------------|-------------------------------------------|--------------------------------|
| 1       | 510(k) Submission | USFDA Medical Device Portal: https://www.fda.gov/medical-devices | Dossier Preparation and Submission |



Ensure that appropriate classification and regulatory strategy is mapped based on device risk, region, and market launch plan. Cite ISO 13485 and 21 CFR Part 820 for QMS alignment.
"""


    else:
        return f"Generate appropriate Design Output content for section: '{section}' for a device named '{device_name}' with intended use '{intended_use}'."

# --- /generate Design Input ---
@app.post("/generate")
async def generate_response(data: DeviceRequest):
    outputs = {}
    for section in data.sections:
        prompt = generate_prompt(data.deviceName, data.intendedUse, section)
        completion = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        outputs[section] = completion.choices[0].message.content.strip()

    return {"results": outputs}

@app.post("/generate-docx")
async def generate_word(data: DeviceRequest):
    doc = WordDoc()

    # Set font globally
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'
    style.font.size = Pt(12)

    section = doc.sections[0]

    # Header
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    header_table.autofit = False
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(2)

    # Logo
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_para.add_run().add_picture("meril_logo.jpg", width=Inches(1.1))

    # Title
    center_cell = header_table.cell(0, 1)
    center_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = center_para.add_run("Design Input")
    run.bold = True
    run.font.size = Pt(17)
    run.font.name = 'Helvetica'

    # Doc Number
    right_cell = header_table.cell(0, 2)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = right_para.add_run(f"Document Number: DI/{data.deviceName[:3].upper()}/001\nRev. 00")
    run.font.size = Pt(11)
    run.font.name = 'Helvetica'

    # Line under header
    header_line = header.add_paragraph()
    header_line_format = header_line.paragraph_format
    header_line_format.space_before = Pt(2)
    header_line_format.space_after = Pt(2)
    hr = header_line.add_run("―" * 54)
    hr.font.name = 'Helvetica'
    hr.font.size = Pt(8)

    # Footer
    footer = section.footer
    footer_line = footer.add_paragraph()
    footer_line.add_run("―" * 54).font.size = Pt(8)
    footer_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer_paragraph = footer.add_paragraph()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.add_run("Meril Healthcare Pvt. Ltd.\nConfidential Document - Page ")
    run.font.size = Pt(10)
    run.font.name = 'Helvetica'
    insert_page_number(footer_paragraph)

    # First page: title (no excessive spacing)
    doc.add_paragraph()
    for _ in range(6): doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run(f"Design Input – {data.deviceName}")
    run.bold = True
    run.font.size = Pt(23)
    run.font.name = 'Helvetica'

    # TOC on page 2
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    numbered_sections = [f"{i+1}. {title}" for i, title in enumerate(data.sections)]
    for sec in numbered_sections:
        para = doc.add_paragraph(sec)
        para.style.font.name = 'Helvetica'
        para.paragraph_format.space_after = Pt(4)

    # Prompts
    prompts = [(section, generate_prompt(data.deviceName, data.intendedUse, section)) for section in data.sections]

    async def fetch(section, prompt):
        try:
            response = await asyncio.wait_for(
                openai.ChatCompletion.acreate(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5
                ),
                timeout=30
            )
            raw = response.choices[0].message.content.strip()
            cleaned = re.sub(r"[\*\#]+", "", raw)
            cleaned = re.sub(r"\n(?=\d+\.)", "\n", cleaned)

            # Parse lines and bold subsection titles
            lines = cleaned.split("\n")
            formatted = []
            for line in lines:
                match = re.match(r"^(\d+\.\s*)([A-Z].+)", line)
                if match:
                    formatted.append(("bold", match.group(0)))
                else:
                    formatted.append(("normal", line))
            return section, formatted
        except Exception as e:
            return section, [("normal", f"⚠️ Error generating section: {str(e)}")]

    results = await asyncio.gather(*[fetch(s, p) for s, p in prompts])

    # Section content
    for i, (section, lines) in enumerate(results):
        doc.add_page_break()

        heading = doc.add_paragraph()
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = heading.add_run(f"{i+1}. {section}")
        run.bold = True
        run.font.size = Pt(15)
        run.font.name = 'Helvetica'

        for tag, line in lines:
            if not line.strip():
                continue
            if tag == "bold":
                # Add 1 line space before subsection title
                doc.add_paragraph()

            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Pt(0 if tag == "bold" else 8)
            run = para.add_run(line.strip())
            run.font.name = 'Helvetica'
            run.font.size = Pt(12)
            if tag == "bold":
                run.bold = True

    # Save file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=Design_Input_{data.deviceName.replace(' ', '_')}.docx"
        }
    )

class DOExportRequest(BaseModel):
    deviceName: str
    intendedUse: str
    sections: list[str]
    results: dict

@app.post("/generate-do-docx")
async def generate_do_word(data: DOExportRequest):
    from io import BytesIO

    # Create document and set global font
    doc = WordDoc()
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'
    style.font.size = Pt(12)

    section = doc.sections[0]

    # --- Header ---
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    header_table.autofit = False
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(2)

    # Logo cell
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_para.add_run().add_picture("meril_logo.jpg", width=Inches(1.1))

    # Title cell
    center_cell = header_table.cell(0, 1)
    center_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = center_para.add_run("Design Output")
    run.bold = True
    run.font.size = Pt(17)
    run.font.name = 'Helvetica'

    # Document number cell
    right_cell = header_table.cell(0, 2)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = right_para.add_run(f"Document Number: DO/{data.deviceName[:3].upper()}/001\nRev. 00")
    run.font.size = Pt(11)
    run.font.name = 'Helvetica'

    # Horizontal rule under header
    header_line = header.add_paragraph()
    header_line_format = header_line.paragraph_format
    header_line_format.space_before = Pt(2)
    header_line_format.space_after = Pt(2)
    hr = header_line.add_run("―" * 54)
    hr.font.name = 'Helvetica'
    hr.font.size = Pt(8)

    # --- Footer ---
    footer = section.footer
    footer_line = footer.add_paragraph()
    footer_line.add_run("―" * 54).font.size = Pt(8)
    footer_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer_paragraph = footer.add_paragraph()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.add_run("Meril Healthcare Pvt. Ltd.\nConfidential Document - Page ")
    run.font.size = Pt(10)
    run.font.name = 'Helvetica'
    insert_page_number(footer_paragraph)

    # --- Title Page ---
    doc.add_paragraph()
    for _ in range(6):
        doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run(f"Design Output – {data.deviceName}")
    run.bold = True
    run.font.size = Pt(23)
    run.font.name = 'Helvetica'

    # --- Table of Contents ---
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for i, section_title in enumerate(data.sections):
        para = doc.add_paragraph(f"{i+1}. {section_title}")
        para.paragraph_format.space_after = Pt(4)
        para.style.font.name = 'Helvetica'

    # --- Fetch AI content ---
    prompts = [
        (section, generate_do_prompt(data.deviceName, data.intendedUse, section))
        for section in data.sections
    ]

    async def fetch(section, prompt):
        try:
            response = await asyncio.wait_for(
                openai.ChatCompletion.acreate(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5
                ),
                timeout=30
            )
            raw = response.choices[0].message.content.strip()
            cleaned = re.sub(r"[#\*]+", "", raw)
            lines = cleaned.split("\n")
            formatted = []
            for line in lines:
                if re.match(r"^\d+\.\s+[A-Z]", line.strip()) or re.match(r"^[-•]", line.strip()):
                    formatted.append(("bold", line))
                else:
                    formatted.append(("normal", line))
            return section, formatted
        except Exception as e:
            return section, [("normal", f"⚠️ Error: {str(e)}")]

    results = await asyncio.gather(*[fetch(s, p) for s, p in prompts])

    # --- Insert sections with real tables and spacing fixes ---
    for i, (section_title, formatted_lines) in enumerate(results):
        doc.add_page_break()
        heading = doc.add_paragraph()
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = heading.add_run(f"{i+1}. {section_title}")
        run.bold = True
        run.font.size = Pt(15)
        run.font.name = 'Helvetica'

        idx = 0
        while idx < len(formatted_lines):
            tag, content = formatted_lines[idx]

            # Table detection: lines starting and ending with "|"
            if content.strip().startswith("|") and content.strip().endswith("|"):
                # Gather the full markdown table block
                table_block = []
                while (
                    idx < len(formatted_lines)
                    and formatted_lines[idx][1].strip().startswith("|")
                    and formatted_lines[idx][1].strip().endswith("|")
                ):
                    table_block.append(formatted_lines[idx][1].strip())
                    idx += 1

                # Parse rows and skip the separator row of hyphens
                raw_rows = [row.strip("|").split("|") for row in table_block]
                header_cells = raw_rows[0]
                data_rows = [
                    row for row in raw_rows[1:]
                    if not all(re.fullmatch(r"-+", cell.strip()) for cell in row)
                ]
                rows = [header_cells] + data_rows

                # Create a real docx table
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"

                # Fill header row
                for col_idx, header_text in enumerate(rows[0]):
                    tbl.rows[0].cells[col_idx].text = header_text.strip()

                # Fill data rows
                for row_idx, row_cells in enumerate(rows[1:], start=1):
                    for col_idx, cell_text in enumerate(row_cells):
                        tbl.rows[row_idx].cells[col_idx].text = cell_text.strip()

                # Add a blank line after each table
                doc.add_paragraph()
                continue  # skip the normal paragraph logic

            # Normal paragraph logic
            if content.strip():
                # Add a blank line before each subsection title
                if tag == "bold":
                    doc.add_paragraph()

                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(0 if tag == "bold" else 8)
                run = para.add_run(content)
                run.font.name = 'Helvetica'
                run.font.size = Pt(12)
                if tag == "bold":
                    run.bold = True

            idx += 1

    # --- Save and return the .docx ---
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=Design_Output_{data.deviceName.replace(' ', '_')}.docx"
        }
    )

# --- /generate-do (Design Output) ---
@app.post("/generate-do")
async def generate_design_output(data: DesignOutputRequest):
    prompt = generate_do_prompt(data.deviceName, data.intendedUse, data.section)
    try:
        response = await openai.ChatCompletion.acreate(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return {"result": response.choices[0].message.content.strip()}
    except Exception as e:
        return {"error": str(e)}

# In-memory storage for finalized DI entries
finalized_devices_db = []
DATA_FILE = Path("finalized_data.json")

@app.on_event("startup")
async def load_finalized_data():
    global finalized_devices_db
    if DATA_FILE.exists():
        with open(DATA_FILE, "r") as f:
            finalized_devices_db = json.load(f)

class FinalizedDevice(BaseModel):
    deviceName: str
    intendedUse: str
    designInputHtml: str
    finalizedBy: str
    diComplete: bool
    doComplete: bool
    finalizedAt: str
    sections: list[str] 


@app.post("/finalize-di")
async def save_finalized_di(data: FinalizedDevice):
    record = data.dict()
    finalized_devices_db.insert(0, record)

    # Save to file
    with open(DATA_FILE, "w") as f:
        json.dump(finalized_devices_db, f, indent=2)

    return {"message": "Saved successfully"}

@app.get("/finalized-devices")
async def get_finalized_devices() -> List[FinalizedDevice]:
    return finalized_devices_db

@app.post("/update-section")
async def update_section(data: UpdateRequest):
    prompt = f"""Revise the following Design Input content for the medical device '{data.deviceName}', intended for '{data.intendedUse}', under the section '{data.section}'.

Only make precise updates based on the user remark provided below. Do not rewrite the entire section. Only modify or remove the specific sentence or subsection as per the remark. Maintain the original structure.

User Remark:
{data.remark}

Current Section Content:
{data.currentContent}
"""
    try:
        response = await openai.ChatCompletion.acreate(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        return {"result": response.choices[0].message.content.strip()}
    except Exception as e:
        return {"error": str(e)}

@app.post("/extract-options")
async def extract_options(payload: dict):
    device_name = payload.get("deviceName", "")
    intended_use = payload.get("intendedUse", "")
    sections = payload.get("sections", [])
    html = payload.get("designInputHtml", "")

    soup = BeautifulSoup(html, "html.parser")
    parsed = {}

    # Knowledge base linking standards to their meanings
    STANDARD_MEANINGS = {
        # Sterilization
        "ISO 11135": "Ethylene Oxide Sterilization",
        "ISO 11137": "Gamma Radiation Sterilization",
        "AAMI TIR28": "EO Sterilization Validation",
        "ISO 17665": "Steam Sterilization",
        # Biocompatibility
        "ISO 10993-5": "Cytotoxicity Testing",
        "ISO 10993-10": "Irritation/Sensitization Testing",
        "USP <87>": "In Vitro Cytotoxicity",
        "USP <88>": "In Vivo Biocompatibility",
        # Packaging
        "ISO 11607": "Packaging Validation",
        "ASTM D4169": "Distribution Simulation Testing",
        # Labeling
        "EN ISO 15223-1": "Medical Device Symbols",
        "21 CFR Part 801": "US Labeling Requirements",
        # Quality Systems
        "ISO 13485": "Quality Management System",
        "21 CFR Part 820": "US FDA QSR"
    }

    SECTION_PROMPTS = {
        "Sterilization Requirements": """
        Analyze the sterilization content and return consolidated options that combine:
        - Methods with their parameters (e.g., "Ethylene Oxide @ 55°C for 12hrs")
        - Standards with their meanings (convert "ISO 11135" to "Ethylene Oxide Sterilization (ISO 11135)")
        - Critical parameters (e.g., "SAL 10^-6", "Residual limits ≤4mg EO")
        Return only a bulleted list of comprehensive options.
        """,
        
        "Biological and Safety Requirements": """
        Extract and consolidate biocompatibility information:
        - Combine test names with standards (e.g., "Cytotoxicity per ISO 10993-5")
        - Include acceptance criteria when mentioned (e.g., "Grade ≤2 Cytotoxicity")
        Return only a bulleted list of combined items.
        """,
        
        "Packaging and Shipping Requirements": """
        Extract packaging information as combined concepts:
        - Packaging types with materials (e.g., "Tyvek/PE Pouch")
        - Tests with purposes (e.g., "Seal Strength ≥2N per ASTM F88")
        - Environmental conditions if specified
        Return only a bulleted list of comprehensive options.
        """,
        
        # ... (similar consolidated prompts for other sections)
    }

    for section in sections:
        parsed[section] = []
        normalized = section.replace(" ", "").replace("/", "").replace("-", "")
        
        # Extract section content (same as before)
        section_div = soup.find("div", {"id": f"section-block-{normalized}"})
        if not section_div:
            header = soup.find(["h2", "h3"], string=lambda text: section.lower() in text.lower() if text else False)
            section_div = header.find_parent("div") if header else None

        if section_div:
            content_div = section_div.find("div", {"id": f"result-{normalized}"}) or \
                         section_div.find("div", class_=lambda x: x and "results" in x.lower()) or \
                         section_div
            text = content_div.get_text(separator="\n", strip=True)

            if section in SECTION_PROMPTS:
                try:
                    prompt = f"""Device: {device_name}
                    Intended Use: {intended_use}
                    Section: {section}
                    
                    {SECTION_PROMPTS[section]}
                    
                    Content to analyze:
                    {text}
                    """
                    
                    response = await openai.ChatCompletion.acreate(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.2,  # Lower temp for more consistent linking
                        max_tokens=600
                    )
                    
                    # Post-process to ensure standards are properly linked
                    raw_options = response.choices[0].message.content.strip()
                    options = []
                    
                    for line in raw_options.split("\n"):
                        if line.strip().startswith("-"):
                            option = line.strip("- ").strip()
                            # Enhance with standard meanings where applicable
                            for std, meaning in STANDARD_MEANINGS.items():
                                if std in option and meaning not in option:
                                    option = option.replace(std, f"{meaning} ({std})")
                            options.append(option)
                    
                    parsed[section] = sorted(list(set(options)))
                    
                except Exception as e:
                    print(f"Error processing {section}: {str(e)}")
                    parsed[section] = []

    return {"parsed": parsed}
